import streamlit as st
import pandas as pd
import matplotlib.pyplot as plt
import seaborn as sns
import numpy as np
import plotly.express as px
import plotly.graph_objects as go
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from openai import OpenAI
from io import BytesIO
import tempfile
import os
import time
import base64
from datetime import datetime
import statsmodels.api as sm
from scipy import stats
from sklearn.preprocessing import StandardScaler
from sklearn.decomposition import PCA
from PIL import Image


# Define color scales dictionary for consistent use throughout the app
color_scales = {
    "Blues": px.colors.sequential.Blues,
    "Viridis": px.colors.sequential.Viridis,
    "Plasma": px.colors.sequential.Plasma,
    "Inferno": px.colors.sequential.Inferno,
    "Magma": px.colors.sequential.Magma,
    "Cividis": px.colors.sequential.Cividis,
    "Turbo": px.colors.sequential.Turbo,
    "Pastel": px.colors.qualitative.Pastel
}

# Set page configuration
st.set_page_config(
    page_title="Excel Data Analyzer Pro",
    page_icon="📊",
    layout="wide",
    initial_sidebar_state="expanded"
)

# Custom CSS for styling
st.markdown("""
<style>
    .main-header {
        font-size: 2.5rem;
        color: #1E88E5;
        text-align: center;
        margin-bottom: 1rem;
        font-weight: 700;
    }
    .sub-header {
        font-size: 1.5rem;
        color: #0D47A1;
        margin-bottom: 1.5rem;
    }
    .section-header {
        font-size: 1.2rem;
        color: #0277BD;
        margin-top: 1rem;
        margin-bottom: 0.5rem;
    }
    .success-message {
        color: #2E7D32;
        font-weight: 600;
        background-color: #E8F5E9;
        padding: 1rem;
        border-radius: 5px;
        text-align: center;
    }
    .info-box {
        background-color: #E3F2FD;
        padding: 1rem;
        border-radius: 5px;
        margin-bottom: 1rem;
    }
    .stButton button {
        background-color: #1976D2;
        color: white;
        font-weight: 600;
        border-radius: 5px;
        padding: 0.5rem 1rem;
        border: none;
    }
    .stButton button:hover {
        background-color: #1565C0;
    }
    div[data-testid="stExpander"] div[role="button"] p {
        font-size: 1.1rem;
        font-weight: 600;
    }
    .reportview-container .main .block-container {
        padding-top: 2rem;
        padding-bottom: 2rem;
    }
    .stProgress > div > div > div > div {
        background-color: #1976D2;
    }
    .sidebar .sidebar-content {
        background-color: #F5F5F5;
    }
    .css-1d391kg {
        padding: 2rem 1rem;
    }
    .row-widget.stSelectbox > div > div {
        background-color: #F5F5F5;
    }
    footer {
        visibility: hidden;
    }
</style>
""", unsafe_allow_html=True)

# Helper function for session state initialization
def init_session_state():
    if 'processed_data' not in st.session_state:
        st.session_state.processed_data = None
    if 'generate_clicked' not in st.session_state:
        st.session_state.generate_clicked = False
    if 'selected_columns' not in st.session_state:
        st.session_state.selected_columns = []
    if 'report_type' not in st.session_state:
        st.session_state.report_type = "Standard"
    if 'api_key' not in st.session_state:
        st.session_state.api_key = os.getenv("OPENAI_API_KEY", "")  # Initialize with empty string or your API key
    if 'color_theme' not in st.session_state:
        st.session_state.color_theme = "Blues"
    if 'advanced_options' not in st.session_state:
        st.session_state.advanced_options = False
    if 'analysis_results' not in st.session_state:
        st.session_state.analysis_results = {}
    if 'handle_missing' not in st.session_state:
        st.session_state.handle_missing = "Remove"
    if 'outlier_treatment' not in st.session_state:
        st.session_state.outlier_treatment = "None"
    if 'percentile_cap' not in st.session_state:
        st.session_state.percentile_cap = 95
    if 'data_cleaned' not in st.session_state:
        st.session_state.data_cleaned = False
    if 'cleaning_summary' not in st.session_state:
        st.session_state.cleaning_summary = None

# Initialize session state
init_session_state()

# Function to get OpenAI client with API key
def get_openai_client():
    api_key = st.session_state.api_key
    if not api_key:
        return None
    try:
        return OpenAI(api_key=api_key)
    except Exception as e:
        st.error(f"Error initializing OpenAI client: {e}")
        return None

# Define bivariate and multivariate functions
# Bivariate analysis functions
def create_bivariate_visualization(df, x_col, y_col, color_theme):
    # Get the color sequence (default to Blues if not found)
    color_seq = color_scales.get(color_theme, px.colors.sequential.Blues)
    
    # Different chart types based on the data types
    x_is_numeric = df[x_col].dtype.kind in 'ifc'
    y_is_numeric = df[y_col].dtype.kind in 'ifc'
    
    if x_is_numeric and y_is_numeric:
        # Create scatter plot for two numeric variables
        fig = px.scatter(
            df, x=x_col, y=y_col,
            title=f"Relationship between {x_col} and {y_col}",
            color_discrete_sequence=color_seq,
            template="plotly_white",
            trendline="ols"  # Add regression line
        )
        
        # Calculate correlation coefficient
        correlation = df[[x_col, y_col]].corr().iloc[0, 1]
        
        # Add correlation annotation
        fig.add_annotation(
            x=0.5, y=0.95,
            xref="paper", yref="paper",
            text=f"Correlation: {correlation:.3f}",
            showarrow=False,
            font=dict(size=14, color="black"),
            align="center",
            bgcolor="rgba(255, 255, 255, 0.8)",
            bordercolor="black",
            borderwidth=1,
            borderpad=4
        )
        
        # Calculate and add regression equation
        X = sm.add_constant(df[x_col])
        model = sm.OLS(df[y_col], X).fit()
        slope = model.params[1]
        intercept = model.params[0]
        r_squared = model.rsquared
        
        fig.add_annotation(
            x=0.5, y=0.89,
            xref="paper", yref="paper",
            text=f"y = {slope:.3f}x + {intercept:.3f}, R² = {r_squared:.3f}",
            showarrow=False,
            font=dict(size=14, color="black"),
            align="center",
            bgcolor="rgba(255, 255, 255, 0.8)",
            bordercolor="black",
            borderwidth=1,
            borderpad=4
        )
        
    elif x_is_numeric and not y_is_numeric:
        # Box plot (numeric vs categorical)
        fig = px.box(
            df, x=y_col, y=x_col,
            title=f"Distribution of {x_col} by {y_col}",
            color=y_col,
            color_discrete_sequence=color_seq,
            template="plotly_white"
        )
        
        # Add mean points
        means = df.groupby(y_col)[x_col].mean().reset_index()
        fig.add_trace(
            go.Scatter(
                x=means[y_col], 
                y=means[x_col],
                mode='markers',
                marker=dict(
                    color='red',
                    size=10,
                    symbol='x'
                ),
                name='Mean'
            )
        )
        
    elif not x_is_numeric and y_is_numeric:
        # Box plot (categorical vs numeric)
        fig = px.box(
            df, x=x_col, y=y_col,
            title=f"Distribution of {y_col} by {x_col}",
            color=x_col,
            color_discrete_sequence=color_seq,
            template="plotly_white"
        )
        
        # Add mean points
        means = df.groupby(x_col)[y_col].mean().reset_index()
        fig.add_trace(
            go.Scatter(
                x=means[x_col], 
                y=means[y_col],
                mode='markers',
                marker=dict(
                    color='red',
                    size=10,
                    symbol='x'
                ),
                name='Mean'
            )
        )
        
    else:
        # Heatmap for categorical vs categorical
        # Create cross-tabulation
        cross_tab = pd.crosstab(df[y_col], df[x_col], normalize='all') * 100
        
        # Create heatmap
        fig = px.imshow(
            cross_tab,
            title=f"Heatmap of {y_col} vs {x_col} (%)",
            color_continuous_scale=color_scales.get(color_theme, px.colors.sequential.Blues),
            labels=dict(x=x_col, y=y_col, color="Percentage (%)"),
            text_auto='.1f',
            aspect="auto",
            template="plotly_white"
        )
        
        # Add counts as text
        counts = pd.crosstab(df[y_col], df[x_col])
        for i, row in enumerate(cross_tab.index):
            for j, col in enumerate(cross_tab.columns):
                fig.add_annotation(
                    x=col, y=row,
                    text=f"n={counts.iloc[i, j]}",
                    showarrow=False,
                    font=dict(color="black", size=9)
                )
                
    # Update layout for better appearance
    fig.update_layout(
        height=600,
        plot_bgcolor='white',
        font=dict(size=12),
        margin=dict(l=80, r=40, t=80, b=80)
    )
    
    return fig

# Calculate bivariate statistics
def bivariate_stats(df, x_col, y_col):
    x_is_numeric = df[x_col].dtype.kind in 'ifc'
    y_is_numeric = df[y_col].dtype.kind in 'ifc'
    
    results = {}
    
    if x_is_numeric and y_is_numeric:
        # Correlation tests and stats for numeric-numeric
        correlation = df[[x_col, y_col]].corr().iloc[0, 1]
        
        # Calculate regression statistics
        X = sm.add_constant(df[x_col])
        model = sm.OLS(df[y_col], X).fit()
        
        results = {
            "correlation": correlation,
            "correlation_type": "Pearson",
            "r_squared": model.rsquared,
            "p_value": model.f_pvalue,
            "slope": model.params[1],
            "intercept": model.params[0],
            "std_error": model.bse[1]
        }
        
    elif (x_is_numeric and not y_is_numeric) or (not x_is_numeric and y_is_numeric):
        # ANOVA for categorical-numeric
        cat_col = y_col if not y_is_numeric else x_col
        num_col = x_col if not y_is_numeric else y_col
        
        categories = df[cat_col].unique()
        data_by_category = [df[df[cat_col] == cat][num_col].dropna() for cat in categories]
        
        # Run ANOVA test
        f_stat, p_value = stats.f_oneway(*data_by_category)
        
        # Calculate group statistics
        group_stats = df.groupby(cat_col)[num_col].agg(['mean', 'std', 'count']).reset_index()
        
        results = {
            "test_type": "ANOVA",
            "f_statistic": f_stat,
            "p_value": p_value,
            "significant": p_value < 0.05,
            "group_stats": group_stats
        }
        
    else:
        # Chi-square test for categorical-categorical
        from scipy.stats import chi2_contingency
        
        # Create contingency table
        contingency_table = pd.crosstab(df[x_col], df[y_col])
        
        # Perform chi-square test
        chi2, p, dof, expected = chi2_contingency(contingency_table)
        
        # Calculate Cramer's V (effect size)
        n = contingency_table.sum().sum()
        phi2 = chi2 / n
        r, k = contingency_table.shape
        phi2corr = max(0, phi2 - ((k-1)*(r-1))/(n-1))
        rcorr = r - ((r-1)**2)/(n-1)
        kcorr = k - ((k-1)**2)/(n-1)
        cramers_v = np.sqrt(phi2corr / min((kcorr-1), (rcorr-1)))
        
        results = {
            "test_type": "Chi-Square",
            "chi2": chi2,
            "p_value": p,
            "dof": dof,
            "significant": p < 0.05,
            "cramers_v": cramers_v,
            "contingency_table": contingency_table
        }
    
    return results

# Multivariate analysis functions
def create_multivariate_visualization(df, columns, color_theme):

    
    # Filter only numeric columns for certain plots
    numeric_columns = [col for col in columns if df[col].dtype.kind in 'ifc']
    
    # 1. Create correlation heatmap (if we have at least 2 numeric columns)
    if len(numeric_columns) >= 2:
        corr_matrix = df[numeric_columns].corr()
        
        fig1 = px.imshow(
            corr_matrix,
            title="Correlation Matrix Heatmap",
            color_continuous_scale=color_scales.get(color_theme, px.colors.sequential.Blues),
            text_auto='.2f',
            aspect="auto",
            template="plotly_white"
        )
        
        fig1.update_layout(
            height=600,
            width=700,
            plot_bgcolor='white',
            font=dict(size=12)
        )
        
        # 2. Create scatter matrix (if we have at least 2 numeric columns)
        if len(numeric_columns) >= 2:
            # If we have more than 4 numeric columns, limit to avoid excessive plot
            plot_columns = numeric_columns[:4] if len(numeric_columns) > 4 else numeric_columns
            
            # Choose a categorical column for color if available
            categorical_columns = [col for col in columns if col not in numeric_columns]
            color_col = categorical_columns[0] if categorical_columns else None
            
            if color_col:
                fig2 = px.scatter_matrix(
                    df, 
                    dimensions=plot_columns,
                    color=color_col,
                    title="Scatter Matrix with Color by " + color_col,
                    template="plotly_white"
                )
            else:
                fig2 = px.scatter_matrix(
                    df, 
                    dimensions=plot_columns,
                    title="Scatter Matrix",
                    color_discrete_sequence=color_scales.get(color_theme, px.colors.sequential.Blues),
                    template="plotly_white"
                )
            
            fig2.update_layout(
                height=700,
                width=800
            )
            
            # Reduce opacity and size for better visibility
            fig2.update_traces(
                diagonal_visible=False,
                marker=dict(size=5, opacity=0.6)
            )
            
            # 3. Create 3D scatter plot if we have at least 3 numeric columns
            if len(numeric_columns) >= 3:
                # Use first three numeric columns for 3D plot
                x_col, y_col, z_col = numeric_columns[:3]
                
                # Choose a categorical column for color if available
                if color_col:
                    fig3 = px.scatter_3d(
                        df, 
                        x=x_col, 
                        y=y_col, 
                        z=z_col,
                        color=color_col,
                        title=f"3D Scatter Plot: {x_col} vs {y_col} vs {z_col}",
                        template="plotly_white"
                    )
                else:
                    fig3 = px.scatter_3d(
                        df, 
                        x=x_col, 
                        y=y_col, 
                        z=z_col,
                        title=f"3D Scatter Plot: {x_col} vs {y_col} vs {z_col}",
                        color_discrete_sequence=color_scales.get(color_theme, px.colors.sequential.Blues),
                        template="plotly_white"
                    )
                
                fig3.update_layout(
                    height=700,
                    width=800
                )
                
                # Make markers smaller and more transparent
                fig3.update_traces(
                    marker=dict(size=3, opacity=0.7)
                )
                
                return [fig1, fig2, fig3]
            
            return [fig1, fig2]
        
        return [fig1]
    
    # If not enough numeric columns, return message
    return ["Not enough numeric columns for multivariate analysis"]

# Perform principal component analysis (PCA)
def perform_pca(df, columns, n_components=2):
    # Filter only numeric columns
    numeric_columns = [col for col in columns if df[col].dtype.kind in 'ifc']
    
    if len(numeric_columns) < 2:
        return {
            "error": "Need at least 2 numeric columns for PCA",
            "pca_result": None,
            "explained_variance": None,
            "loadings": None
        }
    
    # Prepare the data
    X = df[numeric_columns].dropna()
    
    # Standardize the data
    scaler = StandardScaler()
    X_scaled = scaler.fit_transform(X)
    
    # Perform PCA
    n_components = min(n_components, len(numeric_columns))
    pca = PCA(n_components=n_components)
    principal_components = pca.fit_transform(X_scaled)
    
    # Create DataFrame with principal components
    column_names = [f"PC{i+1}" for i in range(n_components)]
    pca_result = pd.DataFrame(data=principal_components, columns=column_names)
    
    # Add a categorical column if available for coloring
    categorical_columns = [col for col in columns if col not in numeric_columns and not pd.api.types.is_datetime64_any_dtype(df[col])]
    if categorical_columns:
        cat_col = categorical_columns[0]
        try:
            pca_result[cat_col] = df[cat_col].values[:len(pca_result)]
        except Exception as e:
            print(f"Warning: Could not add categorical color column '{cat_col}': {e}")
    
    # Calculate and format the explained variance
    explained_variance = pca.explained_variance_ratio_ * 100
    cum_explained_variance = np.cumsum(explained_variance)
    
    variance_df = pd.DataFrame({
        'Component': [f"PC{i+1}" for i in range(n_components)],
        'Explained Variance (%)': explained_variance,
        'Cumulative Variance (%)': cum_explained_variance
    })
    
    # Calculate component loadings (correlations between variables and components)
    loadings = pd.DataFrame(
        pca.components_.T,
        columns=column_names,
        index=numeric_columns
    )
    
    return {
        "pca_result": pca_result,
        "explained_variance": variance_df,
        "loadings": loadings
    }

# Generate AI description for bivariate analysis
def generate_bivariate_description(x_col, y_col, stats):
    client = get_openai_client()
    if not client:
        return "AI insights not available. Please provide an OpenAI API key in the settings."
    
    try:
        prompt = f"""Analyze the relationship between variables '{x_col}' and '{y_col}' based on the following statistics:
        
Statistical Results:
{stats}

Please provide a concise analysis (3-4 sentences) interpreting these results, explaining the relationship between the variables, and highlighting any significant findings or patterns.
"""
        
        response = client.chat.completions.create(
            model="gpt-3.5-turbo",
            messages=[{"role": "user", "content": prompt}],
            max_tokens=300
        )
        return response.choices[0].message.content
    except Exception as e:
        return f"AI insight generation failed. Error: {str(e)}"

# Generate AI description for multivariate analysis
def generate_multivariate_description(columns, pca_results=None, correlation_matrix=None):
    client = get_openai_client()
    if not client:
        return "AI insights not available. Please provide an OpenAI API key in the settings."
    
    try:
        prompt = f"""Analyze the relationships among the variables {', '.join(columns)} based on the following information:
        
"""
        
        if correlation_matrix is not None:
            prompt += f"\nCorrelation Matrix:\n{correlation_matrix.to_string()}\n"
        
        if pca_results is not None and 'explained_variance' in pca_results:
            prompt += f"\nPCA Results - Explained Variance:\n{pca_results['explained_variance'].to_string()}\n"
            
        if pca_results is not None and 'loadings' in pca_results:
            prompt += f"\nPCA Component Loadings (correlations between variables and principal components):\n{pca_results['loadings'].to_string()}\n"
            
        prompt += "\nPlease provide a concise analysis (4-5 sentences) interpreting these results, explaining the relationships among the variables, identifying potential clusters or patterns, and highlighting which variables contribute most to the overall variance in the data."
        
        response = client.chat.completions.create(
            model="gpt-3.5-turbo",
            messages=[{"role": "user", "content": prompt}],
            max_tokens=400
        )
        return response.choices[0].message.content
    except Exception as e:
        return f"AI insight generation failed. Error: {str(e)}"

# Enhanced function to clean data
def clean_data(df, handle_missing="Remove", outlier_treatment="None", percentile_cap=95):
    original_shape = df.shape
    df_copy = df.copy()  # Work with a copy to avoid modifying the original
    
    # Remove completely empty rows and columns
    df_copy = df_copy.dropna(axis=1, how='all')
    df_copy = df_copy.dropna(axis=0, how='all')
    df_copy = df_copy.drop_duplicates()
    
    # Handle missing values based on selection
    if handle_missing == "Remove":
        df_copy = df_copy.dropna()
    elif handle_missing == "Fill with Mean/Mode":
        for col in df_copy.columns:
            if df_copy[col].dtype.kind in 'ifc':  # numeric columns
                df_copy[col].fillna(df_copy[col].mean(), inplace=True)
            else:  # categorical/object columns
                mode_value = df_copy[col].mode()[0] if not df_copy[col].mode().empty else "Unknown"
                df_copy[col].fillna(mode_value, inplace=True)
    
    # Handle outliers if selected
    if outlier_treatment == "Remove" or outlier_treatment == "Cap at percentiles":
        for col in df_copy.columns:
            if df_copy[col].dtype.kind in 'ifc':  # Only process numeric columns
                Q1 = df_copy[col].quantile(0.25)
                Q3 = df_copy[col].quantile(0.75)
                IQR = Q3 - Q1
                lower_bound = Q1 - 1.5 * IQR
                upper_bound = Q3 + 1.5 * IQR
                
                if outlier_treatment == "Remove":
                    df_copy = df_copy[(df_copy[col] >= lower_bound) & (df_copy[col] <= upper_bound)]
                elif outlier_treatment == "Cap at percentiles":
                    lower_cap = df_copy[col].quantile(0.01)
                    upper_cap = df_copy[col].quantile(percentile_cap/100)
                    df_copy[col] = df_copy[col].clip(lower_cap, upper_cap)
    
    # Provide cleaning summary
    new_shape = df_copy.shape
    records_removed = original_shape[0] - new_shape[0]
    columns_removed = original_shape[1] - new_shape[1]
    
    cleaning_summary = {
        "original_records": original_shape[0],
        "original_columns": original_shape[1],
        "records_removed": records_removed,
        "columns_removed": columns_removed,
        "final_records": new_shape[0],
        "final_columns": new_shape[1]
    }
    
    return df_copy, cleaning_summary

# Function to create frequency table
def frequency_table(df, col):
    if df[col].dtype.kind in 'ifc':  # For numeric columns, create bins
        freq = pd.cut(df[col], bins=10).value_counts(dropna=False).sort_index().reset_index()
        freq.columns = ['Range', 'Count']
    else:
        freq = df[col].value_counts(dropna=False).reset_index()
        freq.columns = [col, 'Count']
    
    freq['Percentage'] = round((freq['Count'] / freq['Count'].sum()) * 100, 2)
    freq['Cumulative %'] = freq['Percentage'].cumsum().round(2)
    
    return freq

# Generate descriptive statistics
def descriptive_stats(df, col):
    if df[col].dtype.kind in 'ifc':  # numeric
        stats = {
            "Mean": df[col].mean(),
            "Median": df[col].median(),
            "Min": df[col].min(),
            "Max": df[col].max(),
            "Standard Deviation": df[col].std(),
            "25th Percentile": df[col].quantile(0.25),
            "75th Percentile": df[col].quantile(0.75)
        }
        return pd.DataFrame([stats])
    else:
        return None

# Create enhanced visualizations using Plotly
def create_visualization(df, col, color_theme):
    # Get the color sequence (default to Blues if not found)
    color_seq = color_scales.get(color_theme, px.colors.sequential.Blues)
    
    if df[col].dtype.kind in 'ifc':  # numeric column
        # Create histogram with KDE
        fig = px.histogram(
            df, x=col, 
            marginal="box", 
            title=f"Distribution of {col}",
            color_discrete_sequence=color_seq,
            template="plotly_white"
        )
        
        fig.update_layout(
            xaxis_title=col,
            yaxis_title="Frequency",
            bargap=0.1,
            showlegend=False,
            plot_bgcolor='white',
            height=500
        )
        
        # Add mean line
        fig.add_vline(x=df[col].mean(), line_dash="dash", line_color="red",
                     annotation_text="Mean", annotation_position="top right")
        
        # Add median line
        fig.add_vline(x=df[col].median(), line_dash="dot", line_color="green",
                     annotation_text="Median", annotation_position="top left")
        
    else:  # categorical column
        # Get top 10 categories by frequency
        value_counts = df[col].value_counts().nlargest(10)
        
        # Create a horizontal bar chart
        fig = px.bar(
            x=value_counts.values,
            y=value_counts.index,
            orientation='h',
            title=f"Top 10 Categories in {col}",
            color=value_counts.values,
            color_continuous_scale=color_scales.get(color_theme, px.colors.sequential.Blues),
            template="plotly_white"
        )
        
        fig.update_layout(
            xaxis_title="Count",
            yaxis_title=col,
            yaxis={'categoryorder':'total ascending'},
            height=500,
            plot_bgcolor='white'
        )
        
        # Add percentage labels
        total = sum(value_counts.values)
        percentages = [f"{(val/total*100):.1f}%" for val in value_counts.values]
        
        for i, (val, pct) in enumerate(zip(value_counts.values, percentages)):
            fig.add_annotation(
                x=val,
                y=i,
                text=pct,
                showarrow=False,
                xshift=10,
                font=dict(color="black")
            )
    
    return fig

# Function to generate AI description
def generate_description(column_name, freq_table, stats_table=None):
    client = get_openai_client()
    if not client:
        return "AI insights not available. Please provide an OpenAI API key in the settings."
    
    try:
        prompt = f"""Analyze the following data for column '{column_name}':
        
Frequency Table:
{freq_table.to_string(index=False)}

"""
        if stats_table is not None:
            prompt += f"\nDescriptive Statistics:\n{stats_table.to_string(index=False)}\n"
        
        prompt += "\nPlease provide a concise analysis (3-4 sentences) highlighting the key insights, patterns, and any notable observations from this data."
        
        response = client.chat.completions.create(
            model="gpt-3.5-turbo",
            messages=[{"role": "user", "content": prompt}],
            max_tokens=300
        )
        return response.choices[0].message.content
    except Exception as e:
        return f"AI insight generation failed. Error: {str(e)}"

# Enhanced table styling for Word documents
def set_table_borders(table, header_color="#4472C4"):
    tbl = table._tbl
    tblPr = tbl.tblPr or OxmlElement('w:tblPr')
    tbl.append(tblPr)
    
    # Add table style
    tblStyle = OxmlElement('w:tblStyle')
    tblStyle.set(qn('w:val'), 'TableGrid')
    tblPr.append(tblStyle)
    
    # Add table borders
    borders = OxmlElement('w:tblBorders')
    
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), '000000')
        borders.append(border)
    
    tblPr.append(borders)
    
    # Style the header row
    for i, cell in enumerate(table.rows[0].cells):
        shading_elm = OxmlElement('w:shd')
        shading_elm.set(qn('w:fill'), header_color.lstrip('#'))
        cell._tc.get_or_add_tcPr().append(shading_elm)
        
        # Make text white and bold in header
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.font.bold = True

# Function to create and download a Word report
def generate_word_report(df, columns_to_analyze, analysis_results, report_type, color_theme):
    doc = Document()
    
    # Set document properties
    core_properties = doc.core_properties
    core_properties.title = "Data Analysis Report"
    core_properties.author = "Excel Data Analyzer Pro"
    core_properties.created = datetime.now()
    
    # Title Page
    title_paragraph = doc.add_paragraph()
    title_run = title_paragraph.add_run("Data Analysis Report")
    title_run.font.size = Pt(28)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(0, 70, 127)
    title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add date
    date_paragraph = doc.add_paragraph()
    date_run = date_paragraph.add_run(f"Generated on: {datetime.now().strftime('%B %d, %Y')}")
    date_run.font.size = Pt(12)
    date_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_page_break()
    
    # Table of Contents
    doc.add_heading("Table of Contents", level=1)
    doc.add_paragraph("Executive Summary")
    for col in columns_to_analyze:
        doc.add_paragraph(f"Analysis for: {col}")
    
    doc.add_page_break()
    
    # Executive Summary
    doc.add_heading("Executive Summary", level=1)
    summary_text = "This report provides a comprehensive analysis of the uploaded dataset. "
    
    if report_type == "Executive Summary":
        summary_text += "The executive summary highlights key findings and high-level insights. "
    elif report_type == "Detailed Analysis":
        summary_text += "The detailed analysis includes in-depth statistical measures and comprehensive visualizations. "
    
    summary_text += f"A total of {len(columns_to_analyze)} variables were analyzed, revealing patterns and distributions as detailed in the following sections."
    
    doc.add_paragraph(summary_text)
    
    # Data Overview Table
    doc.add_heading("Data Overview", level=2)
    records_count = len(df)
    columns_count = len(df.columns)
    
    overview_table = doc.add_table(rows=1, cols=2)
    overview_table.style = 'Table Grid'
    hdr_cells = overview_table.rows[0].cells
    hdr_cells[0].text = "Metric"
    hdr_cells[1].text = "Value"
    
    row_cells = overview_table.add_row().cells
    row_cells[0].text = "Number of Records"
    row_cells[1].text = str(records_count)
    
    row_cells = overview_table.add_row().cells
    row_cells[0].text = "Number of Variables"
    row_cells[1].text = str(columns_count)
    
    set_table_borders(overview_table)
    
    doc.add_page_break()
    
    # Individual Column Analysis
    for col in columns_to_analyze:
        doc.add_heading(f"Analysis for: {col}", level=1)
        
        # Add descriptive paragraph
        if col in analysis_results and 'description' in analysis_results[col]:
            doc.add_paragraph(analysis_results[col]['description'])
        
        # Add descriptive statistics if available
        if col in analysis_results and 'stats' in analysis_results[col] and analysis_results[col]['stats'] is not None:
            doc.add_heading("Descriptive Statistics", level=2)
            stats_df = analysis_results[col]['stats']
            
            # Create table for statistics
            stats_table = doc.add_table(rows=1, cols=len(stats_df.columns))
            
            # Add headers
            for i, column_name in enumerate(stats_df.columns):
                stats_table.cell(0, i).text = column_name
            
            # Add data
            for _, row in stats_df.iterrows():
                cells = stats_table.add_row().cells
                for i, val in enumerate(row):
                    cells[i].text = f"{val:.2f}" if isinstance(val, (int, float)) else str(val)
            
            set_table_borders(stats_table)
        
        # Add frequency table
        if col in analysis_results and 'freq' in analysis_results[col]:
            doc.add_heading("Frequency Distribution", level=2)
            freq_df = analysis_results[col]['freq']
            
            freq_table = doc.add_table(rows=1, cols=len(freq_df.columns))
            
            # Add headers
            for i, column_name in enumerate(freq_df.columns):
                freq_table.cell(0, i).text = column_name
            
            # Add data
            for _, row in freq_df.iterrows():
                cells = freq_table.add_row().cells
                for i, val in enumerate(row):
                    cells[i].text = f"{val:.2f}" if isinstance(val, (int, float)) else str(val)
            
            set_table_borders(freq_table)
        
        # Add visualization
        if col in analysis_results and 'plot_path' in analysis_results[col]:
            doc.add_heading("Visualization", level=2)
            doc.add_picture(analysis_results[col]['plot_path'], width=Inches(6))

                # Bivariate Analysis Section
        if 'bivariate_insights' in analysis_results and analysis_results['bivariate_insights']:
            doc.add_heading("Bivariate Analysis", level=1)
            for pair, insight in analysis_results['bivariate_insights'].items():
                doc.add_heading(pair, level=2)
                doc.add_paragraph(insight)
            doc.add_page_break()

                # Add bivariate visualizations
        if 'bivariate_plots' in analysis_results:
            for title, path in analysis_results['bivariate_plots'].items():
                doc.add_heading(f"Bivariate Plot: {title}", level=2)
                doc.add_picture(path, width=Inches(6))
            doc.add_page_break()
        
                # Multivariate Analysis Section
        if 'multivariate_insights' in analysis_results and analysis_results['multivariate_insights']:
            doc.add_heading("Multivariate Analysis", level=1)
            for section, insight in analysis_results['multivariate_insights'].items():
                doc.add_heading(section, level=2)
                doc.add_paragraph(insight)
            doc.add_page_break()

                # Add multivariate visualizations
        if 'multivariate_plots' in analysis_results:
            doc.add_heading("Multivariate Visualizations", level=1)
            for i, path in enumerate(analysis_results['multivariate_plots']):
                doc.add_heading(f"Plot {i+1}", level=2)
                doc.add_picture(path, width=Inches(6))
            doc.add_page_break()

        doc.add_page_break()
    
    # Generate an output stream
    output = BytesIO()
    doc.save(output)
    output.seek(0)
    
    return output

# App Header and Description
st.markdown('<h1 class="main-header">Excel Data Analyzer & Report Generator Pro</h1>', unsafe_allow_html=True)

# Sidebar for configurations
with st.sidebar:
    st.markdown('<h3 class="sub-header">Configuration Settings</h3>', unsafe_allow_html=True)
    
    
    st.session_state.color_theme = st.selectbox(
        "Visualization Theme",
        ["Blues", "Viridis", "Plasma", "Inferno", "Magma", "Cividis", "Turbo", "Pastel"]
    )
    
    advanced_options = st.checkbox("Show Advanced Options", value=st.session_state.advanced_options)
    st.session_state.advanced_options = advanced_options
    
    # Data cleaning options
    if st.session_state.advanced_options:
        st.markdown('<div class="section-header">Data Cleaning Options</div>', unsafe_allow_html=True)
        # Use key parameter but don't directly assign to session_state
        handle_missing = st.selectbox(
            "Handle Missing Values",
            ["Remove", "Fill with Mean/Mode", "Keep as is"],
            index=["Remove", "Fill with Mean/Mode", "Keep as is"].index(st.session_state.handle_missing),
            key="handle_missing_select"
        )
        
        outlier_treatment = st.selectbox(
            "Outlier Treatment",
            ["None", "Remove", "Cap at percentiles"],
            index=["None", "Remove", "Cap at percentiles"].index(st.session_state.outlier_treatment),
            key="outlier_treatment_select"
        )
        
        if outlier_treatment == "Cap at percentiles":
            percentile_cap = st.slider("Percentile Cap", 90, 99, int(st.session_state.percentile_cap), key="percentile_cap_slider")
        else:
            percentile_cap = 95
    else:
        # Default values when advanced options are hidden
        handle_missing = "Remove"
        outlier_treatment = "None"
        percentile_cap = 95
    
    # About section in sidebar
    st.markdown("---")
    st.markdown("""
    <div class="info-box">
    <b>About This App</b><br>
    This application allows you to analyze Excel data files, generate visualizations, and create professional reports enhanced with AI insights.
    </div>
                
     ❤️ Support via Google Pay (UPI)
    Scan the QR code below to contribute and help recharge the OpenAI account for AI-powered insight generation::
    """, unsafe_allow_html=True)
      # Load your uploaded QR code
    qr_image = Image.open("qr.jpg")
    st.image(qr_image, caption="Pay via GPay UPI to thakeemeet@oksbi", width=300)
        ## UPGRADE SECTION
    st.markdown("### 🚀 Want Advanced Features?")
    st.write("""
    Upgrade to **Pro** or **Enterprise** and unlock:

    - Unlimited report exports  
    - Export to PDF  
    - Your own logo & branding  
    - Advanced AI-powered summaries  
    - Priority support  
    - API & integration access  
    """)

    st.markdown("### 📬 Contact for Upgrade Plans:")
    st.markdown("""
    - ✉️ Email: [thakermeet@yahoo.in](mailto:thakermeet@yahoo.in)  
    - 📱 Phone: +91-9586861846  
    - 🔗 [Connect on LinkedIn](https://www.linkedin.com/in/meetthaker/)
    """)

# Main content
col1, col2 = st.columns([1, 1])

with col1:
    st.markdown('<div class="section-header">Upload Your Data</div>', unsafe_allow_html=True)
    uploaded_file = st.file_uploader("Upload Excel file (.xlsx, .xls) or CSV", type=["xlsx", "xls", "csv"], help="Your data remains secure and is not stored permanently")

    # Option for demo data
    use_demo_data = st.checkbox("Use demo data instead")

# Main processing logic
if uploaded_file is not None or use_demo_data:
    with col2:
        st.markdown('<div class="section-header">Data Preview</div>', unsafe_allow_html=True)
        
        with st.spinner("Loading data..."):
            # Read Excel file or use demo data
            if use_demo_data:
                # Create sample data
                np.random.seed(42)
                data = {
                    'Age': np.random.normal(35, 10, 100).astype(int),
                    'Income': np.random.normal(50000, 15000, 100),
                    'Experience': np.random.normal(10, 5, 100),
                    'Satisfaction': np.random.choice(['Low', 'Medium', 'High'], 100),
                    'Department': np.random.choice(['Sales', 'Marketing', 'IT', 'HR'], 100),
                    'Performance': np.random.normal(7.5, 1.5, 100),
                }
                df = pd.DataFrame(data)
            else:
                # Read uploaded file
                if uploaded_file.name.endswith('.csv'):
                    df = pd.read_csv(uploaded_file)
                else:
                    df = pd.read_excel(uploaded_file)
            
            # Display data preview
            st.dataframe(df.head(5), use_container_width=True)
            
            # Show data info
            st.markdown(f"<div class='info-box'>Dataset has {df.shape[0]} rows and {df.shape[1]} columns</div>", unsafe_allow_html=True)
        
        # Keep track of previous cleaning options
        if 'previous_cleaning_options' not in st.session_state:
            st.session_state.previous_cleaning_options = (
                st.session_state.handle_missing,
                st.session_state.outlier_treatment,
                st.session_state.percentile_cap
            )
        
        # Add an explicit "Apply Cleaning Options" button
        if st.button("Apply Cleaning Options"):
            # Update session state with current selections
            st.session_state.handle_missing = handle_missing
            st.session_state.outlier_treatment = outlier_treatment
            st.session_state.percentile_cap = percentile_cap
            st.session_state.data_cleaned = False  # Force recleaning
        
        # Check if cleaning options have changed
        current_cleaning_options = (
            st.session_state.handle_missing,
            st.session_state.outlier_treatment,
            st.session_state.percentile_cap
        )
        
        if current_cleaning_options != st.session_state.previous_cleaning_options:
            # Reset analysis results when cleaning options change
            st.session_state.analysis_results = {}
            st.session_state.generate_clicked = False
            st.session_state.data_cleaned = False
            st.session_state.previous_cleaning_options = current_cleaning_options
        
        # Always display cleaning summary section headers
        st.markdown('<div class="section-header">Data Cleaning Summary</div>', unsafe_allow_html=True)
        
        # Clean data only when needed
        if not st.session_state.data_cleaned or st.session_state.processed_data is None:
            with st.spinner("Cleaning data..."):
                cleaned_df, cleaning_summary = clean_data(
                    df, 
                    handle_missing=st.session_state.handle_missing, 
                    outlier_treatment=st.session_state.outlier_treatment, 
                    percentile_cap=st.session_state.percentile_cap
                )
                st.session_state.processed_data = cleaned_df
                st.session_state.data_cleaned = True
                st.session_state.cleaning_summary = cleaning_summary
        else:
            # Use cached cleaning results
            cleaned_df = st.session_state.processed_data
            cleaning_summary = st.session_state.cleaning_summary
        
        # Display cleaning summary metrics (always show this)
        col_a, col_b, col_c = st.columns(3)
        with col_a:
            st.metric("Original Records", cleaning_summary["original_records"])
        with col_b:
            st.metric("Records Removed", cleaning_summary["records_removed"], delta=-cleaning_summary["records_removed"])
        with col_c:
            st.metric("Final Records", cleaning_summary["final_records"])
    
    # Column selection and analysis options
    st.markdown('<div class="section-header">Select Columns to Analyze</div>', unsafe_allow_html=True)
    
    # Add "Select All" option
    all_columns = st.checkbox("Select All Columns", value=False)
    
    if all_columns:
        st.session_state.selected_columns = list(cleaned_df.columns)
    else:
        st.session_state.selected_columns = st.multiselect(
            "Choose columns for analysis",
            options=list(cleaned_df.columns),
            default=list(cleaned_df.columns)[:3] if len(cleaned_df.columns) > 3 else list(cleaned_df.columns)
        )
        # Track changes to selected columns
    if 'previous_univariate_columns' not in st.session_state:
        st.session_state.previous_univariate_columns = st.session_state.selected_columns

    if set(st.session_state.previous_univariate_columns) != set(st.session_state.selected_columns):
        st.session_state.generate_clicked = False
        st.session_state.previous_univariate_columns = st.session_state.selected_columns
    
    # Analysis tabs
    if st.session_state.selected_columns:
        analysis_tab1, analysis_tab2, analysis_tab3 = st.tabs(["Univariate Analysis", "Bivariate Analysis", "Multivariate Analysis"])
        
        with analysis_tab1:
            st.markdown('<div class="section-header">Univariate Analysis</div>', unsafe_allow_html=True)
            
            if not st.session_state.generate_clicked:
                if st.button("Generate Univariate Analysis"):
                    st.session_state.generate_clicked = True
                    
                    # Create a progress bar
                    progress_bar = st.progress(0)
                    status_text = st.empty()
                    
                    # Reset analysis results
                    st.session_state.analysis_results = {}
                    
                    # Calculate total steps for progress bar
                    total_steps = len(st.session_state.selected_columns) * 3  # 3 operations per column
                    completed_steps = 0
                    
                    # Prepare temp directory for plots
                    with tempfile.TemporaryDirectory() as tmpdirname:
                        # Loop through each selected column
                        for column in st.session_state.selected_columns:
                            # Update status
                            status_text.text(f"Analyzing column: {column}")
                            
                            # Calculate frequency table
                            freq = frequency_table(cleaned_df, column)
                            st.session_state.analysis_results[column] = {'freq': freq}
                            completed_steps += 1
                            progress_bar.progress(completed_steps / total_steps)
                            
                            # Calculate descriptive statistics if applicable
                            stats = descriptive_stats(cleaned_df, column)
                            st.session_state.analysis_results[column]['stats'] = stats
                            completed_steps += 1
                            progress_bar.progress(completed_steps / total_steps)
                            
                            # Generate visualization and save to temp file
                            fig = create_visualization(cleaned_df, column, st.session_state.color_theme)
                            plot_path = os.path.join(tmpdirname, f"{column}_plot.png")
                            fig.write_image(plot_path)
                            st.session_state.analysis_results[column]['plot_path'] = plot_path
                            
                            # Generate AI description if API key is provided
                            if st.session_state.api_key:
                                description = generate_description(column, freq, stats)
                                st.session_state.analysis_results[column]['description'] = description
                            
                            completed_steps += 1
                            progress_bar.progress(completed_steps / total_steps)
                        
                        progress_bar.progress(1.0)
                        status_text.empty()
                    
            # Display univariate analysis results
            if st.session_state.generate_clicked and st.session_state.analysis_results:
                # Create tabs for each analyzed column
                tabs = st.tabs(st.session_state.selected_columns)
                
                for i, column in enumerate(st.session_state.selected_columns):
                    with tabs[i]:
                        if column in st.session_state.analysis_results:
                            col1, col2 = st.columns([1, 1])
                            
                            with col1:
                                # Display AI-generated insights if available
                                if 'description' in st.session_state.analysis_results[column]:
                                    st.markdown(f"<div class='info-box'><b>AI Insights:</b><br>{st.session_state.analysis_results[column]['description']}</div>", unsafe_allow_html=True)
                                
                                # Display frequency table
                                st.markdown("<b>Frequency Distribution</b>", unsafe_allow_html=True)
                                st.dataframe(st.session_state.analysis_results[column]['freq'], use_container_width=True)
                                
                                # Display descriptive statistics if available
                                if 'stats' in st.session_state.analysis_results[column] and st.session_state.analysis_results[column]['stats'] is not None:
                                    st.markdown("<b>Descriptive Statistics</b>", unsafe_allow_html=True)
                                    st.dataframe(st.session_state.analysis_results[column]['stats'], use_container_width=True)
                            
                            with col2:
                                # Re-create visualization (don't use saved image)
                                st.markdown("<b>Visualization</b>", unsafe_allow_html=True)
                                fig = create_visualization(cleaned_df, column, st.session_state.color_theme)
                                st.plotly_chart(fig, use_container_width=True)
                    
        with analysis_tab2:
            st.markdown('<div class="section-header">Bivariate Analysis</div>', unsafe_allow_html=True)
            
            col1, col2 = st.columns(2)
            with col1:
                x_column = st.selectbox("Select X Variable:", options=st.session_state.selected_columns, key="x_var")
            with col2:
                y_column = st.selectbox("Select Y Variable:", options=st.session_state.selected_columns, key="y_var")
            
            if st.button("Generate Bivariate Analysis", key="bivariate_btn"):
                with st.spinner("Generating visualization and statistics..."):
                    # Create visualization
                    fig = create_bivariate_visualization(cleaned_df, x_column, y_column, st.session_state.color_theme)
                    # Save plot image
                    with tempfile.NamedTemporaryFile(delete=False, suffix=".png") as tmpfile:
                        fig.write_image(tmpfile.name)
                        if 'bivariate_plots' not in st.session_state.analysis_results:
                            st.session_state.analysis_results['bivariate_plots'] = {}
                        st.session_state.analysis_results['bivariate_plots'][f"{x_column} vs {y_column}"] = tmpfile.name
                    st.plotly_chart(fig, use_container_width=True)
                    
                    # Calculate statistics
                    stats = bivariate_stats(cleaned_df, x_column, y_column)
                    
                    # Display statistics
                    st.markdown("### Statistical Results")
                    
                    if "correlation" in stats:
                        st.markdown(f"**Correlation Coefficient ({stats['correlation_type']}):** {stats['correlation']:.3f}")
                        st.markdown(f"**R-squared:** {stats['r_squared']:.3f}")
                        st.markdown(f"**p-value:** {stats['p_value']:.6f}")
                        st.markdown(f"**Regression Equation:** y = {stats['slope']:.3f}x + {stats['intercept']:.3f}")
                    
                    elif "test_type" in stats and stats["test_type"] == "ANOVA":
                        st.markdown(f"**ANOVA F-statistic:** {stats['f_statistic']:.3f}")
                        st.markdown(f"**p-value:** {stats['p_value']:.6f}")
                        st.markdown(f"**Result:** {'Statistically significant difference between groups' if stats['significant'] else 'No statistically significant difference between groups'}")
                        
                        st.markdown("**Group Statistics:**")
                        st.dataframe(stats["group_stats"], use_container_width=True)
                    
                    elif "test_type" in stats and stats["test_type"] == "Chi-Square":
                        st.markdown(f"**Chi-Square Value:** {stats['chi2']:.3f}")
                        st.markdown(f"**p-value:** {stats['p_value']:.6f}")
                        st.markdown(f"**Degrees of Freedom:** {stats['dof']}")
                        st.markdown(f"**Cramer's V (Effect Size):** {stats['cramers_v']:.3f}")
                        st.markdown(f"**Result:** {'Statistically significant association between variables' if stats['significant'] else 'No statistically significant association between variables'}")
                        
                        st.markdown("**Contingency Table:**")
                        st.dataframe(stats["contingency_table"], use_container_width=True)
                    
                    # Generate AI insight if API key is available
                    if st.session_state.api_key:
                        st.markdown("### AI-Generated Insights")
                        with st.spinner("Generating insights..."):
                            insight = generate_bivariate_description(x_column, y_column, stats)
                            st.info(insight)
                            if 'bivariate_insights' not in st.session_state.analysis_results:
                                st.session_state.analysis_results['bivariate_insights'] = {}
                            st.session_state.analysis_results['bivariate_insights'][f"{x_column} vs {y_column}"] = insight
        
        with analysis_tab3:
            st.markdown('<div class="section-header">Multivariate Analysis</div>', unsafe_allow_html=True)
            
            # Make sure default values are within available options
            numeric_cols = list(cleaned_df.select_dtypes(include=['number']).columns)
            default_cols = numeric_cols[:min(3, len(numeric_cols))] if numeric_cols else []
            # Ensure defaults are in the options list
            default_cols = [col for col in default_cols if col in st.session_state.selected_columns]

            # Initialize session state variables if they don't exist
            if 'multi_select_all' not in st.session_state:
                st.session_state.multi_select_all = False
            if 'previous_multi_columns' not in st.session_state:
                st.session_state.previous_multi_columns = default_cols

            # Function to handle select all checkbox change
            def handle_select_all_change():
                if st.session_state.multi_select_all:
                    st.session_state.previous_multi_columns = list(st.session_state.selected_columns)
                else:
                    # If unchecked, keep the previous selection (could be modified to clear or reset)
                    pass
            
                # Function to handle multiselect change
            def handle_multi_columns_change():
            # If all columns are selected, set select_all to True
                if set(st.session_state.multi_columns) == set(st.session_state.selected_columns) and st.session_state.selected_columns:
                    st.session_state.multi_select_all = True
                # If not all columns are selected but select_all was True, set it to False
                elif st.session_state.multi_select_all:
                    st.session_state.multi_select_all = False
                # Update previous selection
                st.session_state.previous_multi_columns = st.session_state.multi_columns

             # Select All checkbox
            select_all = st.checkbox("Select All Columns for Multivariate Analysis", 
                            value=st.session_state.multi_select_all,
                            key="multi_select_all",
                            on_change=handle_select_all_change)
            
                    # Multi-select widget
            if select_all:
                # If Select All is checked, show all columns as selected but disable the widget
                multi_columns = st.multiselect(
                    "Selected columns for multivariate analysis:",
                    options=st.session_state.selected_columns,
                    default=st.session_state.selected_columns,
                    key="multi_columns",
                    on_change=handle_multi_columns_change
                )
            else:
                # Otherwise, show the regular multiselect with previous selection
                multi_columns = st.multiselect(
                    "Select columns for multivariate analysis (at least 2 columns):",
                    options=st.session_state.selected_columns,
                    default=st.session_state.previous_multi_columns,
                    key="multi_columns",
                    on_change=handle_multi_columns_change)
                
                    # Check if we need to handle any date columns
            date_columns = [col for col in multi_columns if pd.api.types.is_datetime64_any_dtype(cleaned_df[col])]
            
            # Create a working copy of the dataframe to avoid modifying the original
            analysis_df = cleaned_df.copy()
            
            # Convert date columns to numeric representation for analysis
            date_column_mapping = {}  # Keep track of original date columns and their numeric versions
            if date_columns:
                
                for date_col in date_columns:
                    # Convert dates to ordinal values (days since 0001-01-01)
                    # Wrap in try-except to handle potential NaT values
                    try:
                        # First convert to string to avoid KeyError with Timestamp objects
                        analysis_df[f"{date_col}_numeric"] = pd.to_datetime(analysis_df[date_col]).apply(
                            lambda x: x.toordinal() if pd.notnull(x) else None
                        )
                        date_column_mapping[date_col] = f"{date_col}_numeric"
                    except Exception as e:
                        st.warning(f"Could not convert date column '{date_col}' to numeric: {str(e)}")
                        # If conversion fails, we'll exclude this column
            
            # Replace date columns with their numeric versions in analysis_columns
            analysis_columns = []
            for col in multi_columns:
                if col in date_column_mapping:
                    analysis_columns.append(date_column_mapping[col])
                else:
                    analysis_columns.append(col)

            # Count only numeric columns from the selection for PCA purposes
            selected_numeric_cols = [col for col in multi_columns if cleaned_df[col].dtype.kind in 'ifc']
    
            col1, col2 = st.columns(2)
            with col1:
                run_pca = st.checkbox("Perform Principal Component Analysis (PCA)", value=True if len(selected_numeric_cols) >= 2 else False,
                             disabled=len(selected_numeric_cols) < 2)
            with col2:
                if run_pca:
                    # Fix: Handle the case when min and max would be equal
                    min_components = 2
                    max_components = min(5, len(selected_numeric_cols))
                    # If min would equal max, just display the value instead of a slider
                    if min_components == max_components:
                        st.write(f"Number of PCA components: {min_components}")
                        n_components = min_components
                    else:
                        n_components = st.slider("Number of PCA components:", min_value=min_components, max_value=max_components, value=min(2, max_components))
            
            if len(multi_columns) >= 2 and st.button("Generate Multivariate Analysis", key="multivariate_btn"):
                with st.spinner("Generating visualizations and statistics..."):
                    # Create multivariate visualizations
                    figs = create_multivariate_visualization(analysis_df, analysis_columns, st.session_state.color_theme)

                    # Save each multivariate plot as an image and store path
                    if 'multivariate_plots' not in st.session_state.analysis_results:
                        st.session_state.analysis_results['multivariate_plots'] = []

                    import tempfile

                    for i, fig in enumerate(figs):
                        # Only process plotly figures (skip if it's a warning string)
                        if isinstance(fig, go.Figure) or isinstance(fig, px.scatter().__class__):
                            with tempfile.NamedTemporaryFile(delete=False, suffix=".png") as tmpfile:
                                fig.write_image(tmpfile.name)
                                st.session_state.analysis_results['multivariate_plots'].append(tmpfile.name)
                    
                    if isinstance(figs, list) and not isinstance(figs[0], str):
                        for i, fig in enumerate(figs):
                            st.plotly_chart(fig, use_container_width=True)
                    else:
                        st.warning("Not enough numeric columns for multivariate visualization.")
                    
                    # Perform PCA if selected
                    if run_pca and len(selected_numeric_cols) >= 2:
                        numeric_columns = [col for col in multi_columns if cleaned_df[col].dtype.kind in 'ifc']
                        if len(numeric_columns) >= 2:
                            st.markdown("### Principal Component Analysis (PCA)")
                            
                            
                            pca_results = perform_pca(cleaned_df, multi_columns, n_components)
                            
                            if "error" in pca_results and pca_results["error"]:
                                st.warning(pca_results["error"])
                            else:
                                # Display explained variance
                                st.markdown("#### Explained Variance")
                                st.dataframe(pca_results["explained_variance"], use_container_width=True)
                                
                                # Plot explained variance
                                fig_var = px.bar(
                                    pca_results["explained_variance"],
                                    x="Component",
                                    y="Explained Variance (%)",
                                    title="Explained Variance by Principal Components",
                                    color="Explained Variance (%)",
                                    color_continuous_scale=color_scales.get(st.session_state.color_theme, px.colors.sequential.Blues)
                                )
                                
                                fig_var.add_scatter(
                                    x=pca_results["explained_variance"]["Component"],
                                    y=pca_results["explained_variance"]["Cumulative Variance (%)"],
                                    mode="lines+markers",
                                    name="Cumulative Variance (%)",
                                    line=dict(color="red", width=2)
                                )
                                
                                st.plotly_chart(fig_var, use_container_width=True)
                                
                                # Display component loadings
                                st.markdown("#### Component Loadings")
                                st.dataframe(pca_results["loadings"], use_container_width=True)
                                
                                # Create heatmap of loadings
                                fig_load = px.imshow(
                                    pca_results["loadings"],
                                    title="PCA Component Loadings Heatmap",
                                    color_continuous_scale=color_scales.get(st.session_state.color_theme, px.colors.sequential.Blues),
                                    text_auto='.2f',
                                    aspect="auto"
                                )
                                st.plotly_chart(fig_load, use_container_width=True)
                                
                                # Plot PCA result
                                pca_df = pca_results["pca_result"]
                                
                                # 2D PCA Plot
                                if "PC1" in pca_df.columns and "PC2" in pca_df.columns:
                                    categorical_columns = [col for col in multi_columns if col in pca_df.columns and col not in ["PC1", "PC2"]]
                                    color_col = categorical_columns[0] if categorical_columns else None
                                    
                                    if color_col:
                                        fig_pca = px.scatter(
                                            pca_df,
                                            x="PC1",
                                            y="PC2",
                                            color=color_col,
                                            title="PCA: First Two Principal Components",
                                            template="plotly_white"
                                        )
                                    else:
                                        fig_pca = px.scatter(
                                            pca_df,
                                            x="PC1",
                                            y="PC2",
                                            title="PCA: First Two Principal Components",
                                            color_discrete_sequence=color_scales.get(st.session_state.color_theme, px.colors.sequential.Blues),
                                            template="plotly_white"
                                        )
                                    
                                    # Add variable vectors to the plot
                                    loadings = pca_results["loadings"].loc[:, ["PC1", "PC2"]]
                                    
                                    # Scale the loadings for visualization
                                    scaling = 5
                                    
                                    for i, var in enumerate(loadings.index):
                                        fig_pca.add_shape(
                                            type='line',
                                            x0=0, y0=0,
                                            x1=loadings.iloc[i, 0] * scaling,
                                            y1=loadings.iloc[i, 1] * scaling,
                                            line=dict(color='red', width=1),
                                            opacity=0.8
                                        )
                                        
                                        fig_pca.add_annotation(
                                            x=loadings.iloc[i, 0] * scaling * 1.1,
                                            y=loadings.iloc[i, 1] * scaling * 1.1,
                                            text=var,
                                            showarrow=False,
                                            font=dict(color='red', size=10)
                                        )
                                    
                                    st.plotly_chart(fig_pca, use_container_width=True)
                                
                                # 3D PCA Plot if we have at least 3 components
                                if "PC1" in pca_df.columns and "PC2" in pca_df.columns and "PC3" in pca_df.columns:
                                    categorical_columns = [col for col in multi_columns if col in pca_df.columns and col not in ["PC1", "PC2", "PC3"]]
                                    color_col = categorical_columns[0] if categorical_columns else None
                                    
                                    if color_col:
                                        fig_pca3d = px.scatter_3d(
                                            pca_df,
                                            x="PC1",
                                            y="PC2",
                                            z="PC3",
                                            color=color_col,
                                            title="PCA: First Three Principal Components",
                                            template="plotly_white"
                                        )
                                    else:
                                        fig_pca3d = px.scatter_3d(
                                            pca_df,
                                            x="PC1",
                                            y="PC2",
                                            z="PC3",
                                            title="PCA: First Three Principal Components",
                                            color_discrete_sequence=color_scales.get(st.session_state.color_theme, px.colors.sequential.Blues),
                                            template="plotly_white"
                                        )
                                    
                                    fig_pca3d.update_traces(
                                        marker=dict(size=4, opacity=0.7)
                                    )
                                    
                                    st.plotly_chart(fig_pca3d, use_container_width=True)
                        elif run_pca and len(selected_numeric_cols) < 2:
                                st.warning("At least 2 numeric columns are required for PCA. Please select more numeric columns.")
    
                        
                        # Generate AI insight if API key is available
                        if st.session_state.api_key:
                            st.markdown("### AI-Generated Insights")
                            with st.spinner("Generating multivariate insights..."):
                                # Check for numeric columns
                                numeric_columns = [col for col in multi_columns if cleaned_df[col].dtype.kind in 'ifc']
                                if len(numeric_columns) >= 2:
                                    correlation_matrix = cleaned_df[numeric_columns].corr()
                                    insight = generate_multivariate_description(
                                        multi_columns, 
                                        pca_results if run_pca and len(selected_numeric_cols) >= 2 else None, 
                                        correlation_matrix
                                    )
                                    st.info(insight)
                                    if 'multivariate_insights' not in st.session_state.analysis_results:
                                        st.session_state.analysis_results['multivariate_insights'] = {}
                                    st.session_state.analysis_results['multivariate_insights'][f"Multivariate Analysis ({', '.join(multi_columns)})"] = insight
                                else:
                                    st.warning("At least 2 numeric columns are required for AI-generated multivariate insights.")
    
        # Report Generation Section
        if (
            'analysis_results' in st.session_state 
            and (
            st.session_state.analysis_results.get('bivariate_insights') 
            or st.session_state.analysis_results.get('multivariate_insights') 
            or st.session_state.analysis_results.get('freq')
            or any(
            col in st.session_state.analysis_results and 'freq' in st.session_state.analysis_results[col]
            for col in st.session_state.selected_columns
                 )
                 )
            ):
            st.markdown('<h2 class="section-header">Report Generation</h2>', unsafe_allow_html=True)
            
            with st.expander("Generate and Download Report", expanded=False):
                st.markdown("### Create a comprehensive analysis report")
                
                report_title = st.text_input("Report Title:", "Data Analysis Report")
                
                # Generate Report button
                if st.button("Generate Report"):
                    with st.spinner("Generating comprehensive report..."):
                        # Prepare temp directory for plots
                        with tempfile.TemporaryDirectory() as tmpdirname:
                            # Regenerate plot images for report
                            for column in st.session_state.selected_columns:
                                if column in st.session_state.analysis_results:
                                    fig = create_visualization(cleaned_df, column, st.session_state.color_theme)
                                    plot_path = os.path.join(tmpdirname, f"{column}_plot.png")
                                    fig.write_image(plot_path)
                                    st.session_state.analysis_results[column]['plot_path'] = plot_path
                            
                            # Create and download report
                            report_buffer = generate_word_report(
                                cleaned_df,
                                st.session_state.selected_columns,
                                st.session_state.analysis_results,
                                st.session_state.report_type,
                                st.session_state.color_theme
                            )
                            
                            # Create download button for report
                            st.download_button(
                                label="📄 Download Word Report",
                                data=report_buffer,
                                file_name=f"{report_title.replace(' ', '_')}.docx",
                                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                            )
                            
                            st.success("Report generated successfully!")

        else:
            # Welcome message when no file is uploaded
            st.markdown("""
            <div class="info-box" style="text-align: center; padding: 2rem;">
                <h2>Welcome to Excel Data Analyzer Pro! 📊</h2>
                <p>This tool helps you analyze Excel files and generate professional reports with visualizations and AI-powered insights.</p>
                <p>To get started, upload your Excel file using the uploader on the left.</p>
                <ul style="text-align: left; margin-top: 1rem;">
                    <li>Analyze multiple columns with a single click</li>
                    <li>Generate beautiful visualizations</li>
                    <li>Create professional Word reports</li>
                    <li>Get AI-powered insights (requires OpenAI API key)</li>
                </ul>
            </div>
            """, unsafe_allow_html=True)

# Footer
st.markdown("""
<div style="text-align: center; margin-top: 2rem; padding-top: 1rem; border-top: 1px solid #f0f0f0;">
    <p>Excel Data Analyzer Pro © 2025</p>
</div>
""", unsafe_allow_html=True)

                            

